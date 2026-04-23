"""Generate pixo_features_summary.docx and pixo_features_detailed.docx.

Run from the repo root:
    python docs/build_feature_docs.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent

NAVY = RGBColor(0x1F, 0x3A, 0x8A)
GREEN = RGBColor(0x10, 0x7C, 0x10)
YELLOW = RGBColor(0xB4, 0x7B, 0x00)
RED = RGBColor(0xB8, 0x27, 0x27)
CODE_GREEN = RGBColor(0x0B, 0x5E, 0x07)
DIM = RGBColor(0x64, 0x74, 0x8B)


# ------------------ helpers ------------------

class Doc:
    def __init__(self, title: str):
        self.doc = Document()
        s = self.doc.styles["Normal"]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        self.doc.core_properties.title = title

    # Headings
    def h1(self, text: str):
        p = self.doc.add_heading(text, level=1)
        for r in p.runs:
            r.font.color.rgb = NAVY

    def h2(self, text: str):
        self.doc.add_heading(text, level=2)

    def h3(self, text: str):
        self.doc.add_heading(text, level=3)

    # Text
    def p(self, text: str = ""):
        self.doc.add_paragraph(text)

    def bold_line(self, label: str, rest: str = ""):
        p = self.doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        if rest:
            p.add_run(" " + rest)

    def bullet(self, text: str):
        self.doc.add_paragraph(text, style="List Bullet")

    def code_block(self, text: str):
        """Monospace block, each line as its own paragraph."""
        for line in text.splitlines():
            p = self.doc.add_paragraph()
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = CODE_GREEN

    def inline_code(self, paragraph, text: str):
        r = paragraph.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = CODE_GREEN

    def dim(self, text: str):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.italic = True
        r.font.color.rgb = DIM

    def hr(self):
        """Horizontal rule."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def table(self, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            r = p.add_run(h)
            r.bold = True
        for row_i, row in enumerate(rows, start=1):
            cells = t.rows[row_i].cells
            for ci, val in enumerate(row):
                cells[ci].text = str(val)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    if i < len(row.cells):
                        row.cells[i].width = Inches(w)

    def privacy_table(self, rows: list[tuple[str, str]]):
        """Two-column table with colored badge in first column."""
        t = self.doc.add_table(rows=len(rows), cols=2)
        t.style = "Light Grid Accent 1"
        for i, (badge, desc) in enumerate(rows):
            cells = t.rows[i].cells
            p = cells[0].paragraphs[0]
            r = p.add_run(badge)
            r.bold = True
            if badge == "green":
                r.font.color.rgb = GREEN
            elif badge == "yellow":
                r.font.color.rgb = YELLOW
            elif badge == "red":
                r.font.color.rgb = RED
            cells[1].text = desc

    def save(self, filename: str):
        path = OUT_DIR / filename
        self.doc.save(str(path))
        print(f"Wrote: {path}")


# ====================================================================
# SUMMARY DOC
# ====================================================================

def build_summary():
    d = Doc("pixo — Features Summary")

    d.h1("pixo — Features Summary")
    d.p("Quick reference for everything pixo can do. For the full explanation "
        "of each feature (why it exists, what it does, how to use), see "
        "pixo_features_detailed.docx.")
    d.dim("Current version: v0.3.0")
    d.hr()

    # Commands table
    d.h2("Commands at a glance")
    d.table(
        ["Command", "What it does"],
        [
            ["pixo try", "Zero-setup demo. Auto-picks a model + sample image and opens a browser report."],
            ["pixo list", "List all available models with privacy badges."],
            ["pixo info <model>", "Show detailed info and variants for a single model."],
            ["pixo pull <model>", "Download a model to ~/.pixo/models/."],
            ["pixo run <model> -i <file>", "Run inference on an image or video."],
            ["pixo rm <model>", "Delete a downloaded model."],
            ["pixo pipe \"m1 -> m2\" -i <file>", "Chain multiple models together."],
            ["pixo compare <m1> <m2> -i <file>", "Run N detection models and show where they disagree."],
            ["pixo share [job_id]", "Export a self-contained HTML report for a completed run."],
            ["pixo serve <model>", "Launch a Gradio browser UI for one model."],
            ["pixo ui", "Launch the full local dashboard (FastAPI + React)."],
            ["pixo doctor", "Check hardware and show recommendations."],
            ["pixo optimize <model>", "Convert to ONNX for ~40% faster CPU inference."],
            ["pixo history", "List past jobs (running, completed, paused, failed)."],
            ["pixo resume [job_id]", "Continue a paused or failed job."],
            ["pixo view <job_id>", "Open a job's output folder."],
            ["pixo jobs-clean", "Delete completed job checkpoints."],
            ["pixo setup-cloud", "Connect free Kaggle or Colab GPU accounts."],
            ["pixo cloud-status", "Show status of connected cloud backends."],
            ["pixo env-list", "Show isolated per-model environments."],
            ["pixo env-clean <model>", "Remove a model's isolated environment."],
            ["pixo upgrade", "Update pixo to the latest version."],
            ["pixo guide", "Print the in-terminal usage guide."],
        ],
        col_widths=[2.4, 4.2],
    )
    d.hr()

    # Flags
    d.h2("Flags you'll use most")
    d.table(
        ["Flag", "Command", "What it does"],
        [
            ["--input, -i", "run, pipe, compare", "Path to the input image or video."],
            ["--output, -o", "most", "Where to save the output folder."],
            ["--device, -d", "run, compare, pipe", "Force cpu or cuda."],
            ["--backend, -b", "run", "Force local, kaggle, or colab."],
            ["--airgap", "run", "Block all outbound network calls during the run."],
            ["--low-memory", "run", "Frame-by-frame processing for small-RAM laptops."],
            ["--background", "run", "Lowest OS priority so the laptop stays usable."],
            ["--max-ram, --max-cpu", "run", "Cap resource usage."],
            ["--force", "run", "Skip the pre-run safety check."],
            ["--isolate", "pull, run", "Use a per-model virtual env."],
            ["--prompt, -p", "run, pipe", "Text prompt for grounding_dino (\"person, car\")."],
            ["--task", "run", "Task for florence2 (caption, detect, ocr)."],
            ["--conf", "compare", "Confidence threshold for detections."],
            ["--iou", "compare", "IoU threshold for matching boxes across models."],
        ],
        col_widths=[1.6, 1.8, 3.2],
    )
    d.hr()

    # Models
    d.h2("Available models")
    d.table(
        ["Model", "Task", "Default Size", "Notes"],
        [
            ["yolov8", "Object detection", "6 MB", "Fastest, best for demos and first-time use."],
            ["yolov11", "Object detection", "5 MB", "Newer YOLO, slightly more accurate than v8."],
            ["yolov12", "Object detection", "7 MB", "Attention-based, state-of-the-art accuracy."],
            ["rtdetr", "Object detection", "64 MB", "Transformer detector, no NMS."],
            ["grounding_dino", "Text-prompted detection", "341 MB", "\"Find red cars\" — detects anything by natural language."],
            ["florence2", "Vision-language", "460 MB", "Captioning, OCR, detection, all in one model."],
            ["depth_anything_v2", "Depth estimation", "99 MB", "Monocular depth maps from a single image."],
            ["sam2", "Segmentation", "898 MB", "Pixel-perfect masks. Slow on CPU."],
            ["samurai", "Video tracking + segmentation", "898 MB", "Track any object across video frames."],
        ],
        col_widths=[1.4, 1.7, 1.0, 2.5],
    )
    d.p("All nine models carry a GREEN privacy badge — they run fully offline "
        "after the weights are downloaded once.")
    d.hr()

    # Feature groups
    d.h2("Feature groups")

    d.h3("Core runtime")
    d.bullet("Resource Guardian — pre-run safety check, RAM/CPU/GPU limits, temperature monitoring")
    d.bullet("Plugin system — models defined by modelcard.yaml + run.py, easy to add new ones")
    d.bullet("Standard output format — every run produces results.json, summary.txt, visualizations/, exports/")
    d.bullet("ONNX optimization — pixo optimize converts models for faster CPU inference")
    d.bullet("Checkpointing — auto-save every N frames, Ctrl+C pauses, pixo resume continues")

    d.h3("v0.3 — local-first differentiators")
    d.bullet("pixo try — zero-setup demo for new users")
    d.bullet("pixo share — self-contained HTML reports, no server needed")
    d.bullet("pixo compare — cross-model disagreement browser")
    d.bullet("--airgap — blocks all network calls for the duration of a run")
    d.bullet("Privacy badges — green/yellow/red per model")
    d.bullet("pixo serve — Gradio UI for any model")

    d.h3("Cloud")
    d.bullet("Smart router — estimates local vs cloud time before running")
    d.bullet("Kaggle backend — fully automated, uses free 30 hrs/week")
    d.bullet("Colab backend — notebook generation for free T4 GPU")

    d.h3("Pipelines")
    d.bullet("Model piping — pixo pipe \"grounding_dino -> sam2\" chains models")
    d.bullet("Pre-built templates — detect_and_segment, segment_and_depth")

    d.h3("Isolation")
    d.bullet("Per-model venvs — pixo pull <model> --isolate creates an isolated environment")

    d.h3("Python SDK")
    d.bullet("pixo.list_models(), pixo.pull(), pixo.run(), pixo.doctor(), pixo.pipe() — use pixo from Python scripts and notebooks")
    d.hr()

    # Two UIs
    d.h2("Two web UIs — what each is for")
    d.table(
        ["Property", "pixo ui (dashboard)", "pixo serve <model> (Gradio)"],
        [
            ["Purpose", "Browse/manage everything — models, jobs, hardware", "Run one specific model from the browser"],
            ["Pages", "Multiple", "One"],
            ["Install", "pip install pixo[web]", "pip install pixo[demo]"],
            ["Port", "8420", "7860"],
            ["Best for", "\"What does my pixo install look like?\"", "\"Drag-drop an image and see results\""],
        ],
        col_widths=[1.4, 2.6, 2.6],
    )
    d.hr()

    # Privacy badges
    d.h2("Privacy badges")
    d.p("Every model card declares its privacy posture:")
    d.privacy_table([
        ("green", "Runs fully offline after weights are pulled."),
        ("yellow", "Needs internet for first weight download, then offline."),
        ("red", "Needs runtime internet (API calls, remote services)."),
    ])
    d.p("Combined with --airgap, this gives users an end-to-end \"no bytes "
        "leave my laptop\" guarantee — something no competitor offers as a "
        "first-class feature.")
    d.hr()

    # Output layout
    d.h2("Output structure")
    d.p("Every pixo run produces this folder layout:")
    d.code_block(
        "pixo_output/\n"
        "  <jobid>_<model>_<input>/\n"
        "    results.json       # machine-readable metadata\n"
        "    summary.txt        # human-readable summary\n"
        "    visualizations/    # annotated images or video\n"
        "    raw/               # raw model output files\n"
        "    exports/           # COCO JSON, CSV, etc."
    )
    d.p("pixo share <job_id> turns this folder into a single shareable .html file.")
    d.hr()

    # Resource safety
    d.h2("Resource safety")
    d.p("Every pixo run goes through a pre-run check that looks at:")
    d.bullet("Free RAM vs model size")
    d.bullet("Free disk vs expected output")
    d.bullet("GPU VRAM if using --device cuda")
    d.bullet("CPU temperature")
    d.p("If anything is unsafe, pixo suggests alternatives (smaller variant, "
        "--low-memory, cloud backend) instead of crashing your laptop.")
    d.hr()

    # File locations
    d.h2("File locations")
    d.table(
        ["Path", "What's in it"],
        [
            ["~/.pixo/models/", "Downloaded model weights"],
            ["~/.pixo/envs/", "Isolated model environments"],
            ["~/.pixo/jobs/", "Job checkpoints for resume"],
            ["~/.pixo/shares/", "Self-contained HTML share bundles"],
            ["~/.pixo/samples/", "Cached sample images for pixo try"],
            ["~/.pixo/config.yaml", "Cloud backend credentials"],
            ["~/.pixo/logs/", "Run logs"],
            ["./pixo_output/", "Default output for runs"],
        ],
        col_widths=[2.2, 4.2],
    )

    d.save("pixo_features_summary.docx")


# ====================================================================
# DETAILED DOC
# ====================================================================

def build_detailed():
    d = Doc("pixo — Features Explained")

    d.h1("pixo — Features Explained")
    d.p("Every feature in pixo, explained in plain English.")
    d.p("For each feature: why it exists, what it does, how to use it. "
        "For a one-page overview instead, see pixo_features_summary.docx.")
    d.dim("Current version: v0.3.0")
    d.hr()

    # TOC
    d.h2("Table of contents")
    for i, t in enumerate([
        "Getting started — zero-setup demo",
        "Running models",
        "Seeing what's available",
        "The nine bundled models",
        "Privacy and airgap mode",
        "Sharing your results",
        "Comparing models",
        "Browser UIs",
        "Resource safety — not freezing your laptop",
        "Checkpointing — never lose progress",
        "Standard output structure",
        "Free cloud GPUs",
        "ONNX optimization",
        "Isolated environments",
        "Model piping",
        "Python SDK",
        "Hardware profiling",
        "Everyday housekeeping",
    ], start=1):
        d.bullet(f"{i}. {t}")
    d.hr()

    # ---------- 1 ----------
    d.h2("1. Getting started — zero-setup demo")
    d.h3("pixo try")
    d.bold_line("Why it exists")
    d.p("Most people who try a new CLI give up within 60 seconds if they "
        "can't see it work. Every step between \"I just installed this\" and "
        "\"I see output\" is a chance to bounce. pixo try exists to collapse "
        "those steps into one command.")
    d.bold_line("What it does")
    d.bullet("Detects your hardware and picks a suitable model (YOLOv8 on CPU laptops, YOLOv11 if you have a GPU).")
    d.bullet("Finds a sample image — first checks if Ultralytics bundled one, then falls back to ~/.pixo/samples/, then downloads if needed.")
    d.bullet("Auto-downloads the chosen model if it isn't installed yet.")
    d.bullet("Runs the model on the sample image.")
    d.bullet("Generates a self-contained HTML report.")
    d.bullet("Opens that report in your default browser.")
    d.bold_line("How to use it")
    d.code_block("pixo try")
    d.p("That's it. No flags required. Overrides:")
    d.code_block(
        "pixo try --model yolov11            # pick a specific model\n"
        "pixo try --input photo.jpg          # use your own image\n"
        "pixo try --model sam2 --input cat.png"
    )
    d.bold_line("Who it's for")
    d.p("A brand-new user who just ran pip install pixo and wants to see it "
        "work. Not meant for day-to-day use — once you know pixo, you'll "
        "use pixo run directly.")
    d.hr()

    # ---------- 2 ----------
    d.h2("2. Running models")
    d.h3("pixo run")
    d.bold_line("Why it exists")
    d.p("This is the main command. Everything else supports it.")
    d.bold_line("What it does")
    d.p("Runs any model in pixo's catalog on an image or video, handles all "
        "the plumbing (downloading weights if needed, picking the right "
        "device, collecting output), and saves results in a standard folder "
        "layout.")
    d.bold_line("How to use it — common patterns")
    d.code_block(
        "# Basic detection on an image\n"
        "pixo run yolov8 --input photo.jpg\n\n"
        "# Detection on a video\n"
        "pixo run yolov8 --input video.mp4\n\n"
        "# Use a specific variant\n"
        "pixo run yolov8:small --input photo.jpg\n\n"
        "# Use your GPU explicitly\n"
        "pixo run yolov8 --input photo.jpg --device cuda\n\n"
        "# Save output somewhere specific\n"
        "pixo run yolov8 --input photo.jpg --output ~/Desktop/results/\n\n"
        "# Text-prompted detection\n"
        "pixo run grounding_dino --input photo.jpg --prompt \"red car, yellow hat\"\n\n"
        "# Multi-task vision-language model\n"
        "pixo run florence2 --input photo.jpg --task caption\n"
        "pixo run florence2 --input photo.jpg --task ocr\n"
        "pixo run florence2 --input photo.jpg --task detect\n\n"
        "# Low-memory laptop-safe mode\n"
        "pixo run yolov8 --input video.mp4 --low-memory\n\n"
        "# Run in the background\n"
        "pixo run yolov8 --input video.mp4 --background\n\n"
        "# Run with network blocked\n"
        "pixo run yolov8 --input photo.jpg --airgap\n\n"
        "# Force a cloud backend\n"
        "pixo run sam2 --input photo.jpg --backend kaggle"
    )
    d.bold_line("What happens under the hood")
    d.bullet("Checks the input file exists.")
    d.bullet("Checks if the model is downloaded. If not, pulls from HuggingFace.")
    d.bullet("Runs a resource safety check (RAM, disk, CPU temperature).")
    d.bullet("Picks a backend — local, or a cloud GPU if configured.")
    d.bullet("Uses an ONNX-optimized version if pixo optimize was run earlier.")
    d.bullet("Loads the model and runs it, yielding progress updates.")
    d.bullet("Saves results in the standard output format.")
    d.bullet("If interrupted (Ctrl+C), pauses gracefully and saves a checkpoint.")
    d.hr()

    # ---------- 3 ----------
    d.h2("3. Seeing what's available")
    d.h3("pixo list")
    d.p("Shows every model pixo knows about as a table with download status, "
        "name, task, variants, default size, privacy badge, and a short "
        "description.")
    d.code_block("pixo list")
    d.h3("pixo info <model>")
    d.p("Detailed info about one model: author, source repo, input/output "
        "types, privacy posture, all variants with sizes, minimum RAM, and "
        "whether the runner is ready.")
    d.code_block("pixo info yolov8\npixo info sam2")
    d.hr()

    # ---------- 4 ----------
    d.h2("4. The nine bundled models")

    d.h3("YOLO family — fast object detection")
    d.p("Three generations of the most popular real-time detector.")
    d.bullet("yolov8 — the workhorse. Fast on CPU (~7s for a photo), very accurate. Default choice.")
    d.bullet("yolov11 — newer architecture, slightly better accuracy.")
    d.bullet("yolov12 — latest, attention-based, state-of-the-art.")
    d.p("All three detect the standard 80 COCO classes (person, car, dog, chair, etc.).")
    d.code_block("pixo run yolov8 --input photo.jpg")

    d.h3("RT-DETR — transformer detector")
    d.p("rtdetr uses a transformer architecture and doesn't need post-processing "
        "(no NMS step). More accurate than YOLO in some scenarios but a larger model.")
    d.code_block("pixo run rtdetr --input photo.jpg")

    d.h3("Grounding DINO — text-prompted detection")
    d.p("grounding_dino is different from YOLO. Instead of detecting from a fixed "
        "list of classes, you give it natural-language prompts and it finds whatever "
        "you describe.")
    d.code_block(
        "pixo run grounding_dino --input photo.jpg --prompt \"red backpack\"\n"
        "pixo run grounding_dino --input photo.jpg --prompt \"person wearing a hat\""
    )
    d.p("Use it when the thing you want to detect isn't in YOLO's class list.")

    d.h3("Florence-2 — vision-language model")
    d.p("florence2 is a multi-task model: one weights file, multiple things it can do.")
    d.code_block(
        "pixo run florence2 --input photo.jpg --task caption           # describe\n"
        "pixo run florence2 --input photo.jpg --task detailed_caption  # longer description\n"
        "pixo run florence2 --input photo.jpg --task detect            # find objects\n"
        "pixo run florence2 --input photo.jpg --task ocr               # extract text"
    )

    d.h3("Depth Anything v2 — depth estimation")
    d.p("depth_anything_v2 produces a depth map from a single photo — no stereo "
        "camera or special hardware needed. Useful for AR, 3D reconstruction, "
        "robotics, or just for looking at in awe.")
    d.code_block("pixo run depth_anything_v2 --input photo.jpg")

    d.h3("SAM2 — segment anything")
    d.p("sam2 produces pixel-perfect segmentation masks. It's slow on CPU "
        "(several minutes per image) but the output quality is incredible. "
        "If you have a GPU or Kaggle access, this is the gold standard.")
    d.code_block("pixo run sam2 --input photo.jpg --backend kaggle   # recommended for CPU users")

    d.h3("SAMURAI — video tracking")
    d.p("samurai tracks any object across video frames using SAM2's masking. "
        "You give it a video, it produces a tracked, masked output video.")
    d.code_block("pixo run samurai --input video.mp4 --backend kaggle")
    d.hr()

    # ---------- 5 ----------
    d.h2("5. Privacy and airgap mode")
    d.h3("Privacy badges")
    d.bold_line("Why it exists")
    d.p("Most CV tools silently phone home — telemetry, version checks, "
        "HuggingFace lookups, automatic updates. For users handling sensitive "
        "data (medical, legal, journalistic, enterprise), that's a dealbreaker. "
        "Privacy badges give an honest declaration of what each model does.")
    d.bold_line("What it does")
    d.p("Every model card carries one of three badges:")
    d.privacy_table([
        ("green", "Runs fully offline once weights are downloaded."),
        ("yellow", "Needs internet for first weight download but runs offline after."),
        ("red", "Requires runtime internet access (API calls, remote services)."),
    ])
    d.bold_line("How to use it")
    d.code_block("pixo list        # see badges for all models\npixo info sam2   # see the badge plus any notes")
    d.p("All nine bundled models are green.")

    d.h3("--airgap mode")
    d.bold_line("Why it exists")
    d.p("A declared badge is a promise. --airgap is how you verify the promise. "
        "When you need to guarantee that a particular run never reaches the "
        "internet — medical images, legal documents, confidential photos — "
        "you add one flag.")
    d.bold_line("What it does")
    d.bullet("Sets environment variables (HF_HUB_OFFLINE=1, TRANSFORMERS_OFFLINE=1, YOLO_OFFLINE=True) so libraries skip their connectivity checks.")
    d.bullet("Monkey-patches ultralytics.utils.checks.check_online so YOLO doesn't ping 1.1.1.1 before running.")
    d.bullet("Monkey-patches Python's socket and DNS functions so any outbound connection raises AirgapViolation.")
    d.bullet("Reverts all patches when the run finishes — system network is untouched afterward.")
    d.p("Only loopback (127.0.0.1) stays allowed, so local services keep working.")
    d.bold_line("How to use it")
    d.code_block("pixo run yolov8 --input photo.jpg --airgap")
    d.p("You'll see \"Airgap: network access blocked for this run\" in the "
        "output. If anything inside the run tries to dial out, it'll fail "
        "fast with a clear error telling you which host was blocked.")
    d.bold_line("Combines with")
    d.p("You cannot combine --airgap with --backend kaggle or --backend colab "
        "— those need network by definition. pixo will refuse with a clear error.")
    d.hr()

    # ---------- 6 ----------
    d.h2("6. Sharing your results")
    d.h3("pixo share")
    d.bold_line("Why it exists")
    d.p("When you want to show a result to someone — a colleague, on Twitter, "
        "in a pull request — you need a format that \"just works.\" Screenshots "
        "lose metadata. Zip files require unzipping. Web apps need a server. "
        "A single self-contained HTML file is the lowest-friction artifact that "
        "exists: it opens in any browser, anywhere, forever, with no dependencies.")
    d.bold_line("What it does")
    d.p("Takes a completed pixo run job and produces a single .html file with:")
    d.bullet("All thumbnails and visualizations embedded as base64 data URIs.")
    d.bullet("Summary of what the run did (model, device, time, objects found).")
    d.bullet("Provenance footer (pixo version, job ID, timestamp).")
    d.bullet("Dark-mode-aware, responsive styling.")
    d.p("The resulting file opens in any browser with zero network calls.")
    d.bold_line("How to use it")
    d.code_block(
        "pixo share                 # most recent completed run\n"
        "pixo share 3a37cea3        # specific job\n"
        "pixo share --no-open       # don't open browser automatically"
    )
    d.p("The file is saved to ~/.pixo/shares/<job_id>.html. Attach it to a "
        "tweet, Slack message, email, or GitHub issue.")
    d.hr()

    # ---------- 7 ----------
    d.h2("7. Comparing models")
    d.h3("pixo compare")
    d.bold_line("Why it exists")
    d.p("The obvious question a user has with nine detection models is: "
        "\"Which one is best for my image?\" Nobody can answer that "
        "abstractly — it depends on the image. pixo compare lets you see "
        "the actual differences on the actual image, not a benchmark from a paper.")
    d.p("The feature also doubles as something nobody else has shipped as a "
        "flagship experience: most tools show you one model's output. pixo "
        "shows you specifically where the models disagree.")
    d.bold_line("What it does")
    d.bullet("Runs each specified model on the same input image.")
    d.bullet("Collects every detection (box + label + score) from each model.")
    d.bullet("Matches boxes across models using IoU (area overlap).")
    d.bullet("Classifies every object into three groups: Agreement (all models found it), Partial (some did), Only one model (exactly one did).")
    d.bullet("Produces a self-contained HTML report with colored boxes showing which model detected what.")
    d.bold_line("How to use it")
    d.code_block(
        "pixo compare yolov8 yolov11 --input photo.jpg\n"
        "pixo compare yolov8 yolov11 yolov12 --input photo.jpg\n"
        "pixo compare yolov8 yolov11 yolov12 rtdetr --input photo.jpg"
    )
    d.p("Tune the thresholds:")
    d.code_block(
        "pixo compare yolov8 yolov11 --input photo.jpg --conf 0.3\n"
        "pixo compare yolov8 yolov11 --input photo.jpg --iou 0.6"
    )
    d.bold_line("Scope in v0.3")
    d.p("Currently supports yolov8, yolov11, yolov12, rtdetr (all Ultralytics-"
        "backed detection models) on image inputs. Video comparison and "
        "segmentation/depth comparison are planned for v0.4.")
    d.bold_line("Why this is the flagship differentiator")
    d.p("A user picking between YOLOv8 and YOLOv12 can now do so based on "
        "their own photos, not a leaderboard. A researcher writing a paper "
        "gets a clean visualization of disagreement. And it's the kind of "
        "screenshot that goes viral on ML Twitter.")
    d.hr()

    # ---------- 8 ----------
    d.h2("8. Browser UIs")
    d.p("pixo has two browser interfaces that serve different purposes.")

    d.h3("pixo serve <model> — Gradio UI")
    d.bold_line("Why it exists")
    d.p("Sometimes typing commands is slower than dragging a file. Sometimes "
        "you're showing pixo to a non-technical colleague. Sometimes you're "
        "doing a live demo and want something visual.")
    d.bold_line("What it does")
    d.p("Launches a single-page Gradio UI on http://localhost:7860 dedicated "
        "to one model. Drag an image in, click Run, see the annotated output.")
    d.p("For grounding_dino, a text prompt field appears automatically. For "
        "florence2, a task dropdown appears.")
    d.bold_line("How to use it")
    d.code_block(
        "pip install pixo[demo]     # one-time\n"
        "pixo serve yolov8\n"
        "pixo serve grounding_dino\n"
        "pixo serve florence2\n"
        "pixo serve yolov8 --share  # temporary public URL"
    )
    d.bold_line("Best use cases")
    d.bullet("Showing pixo at a meetup or on a screenshare.")
    d.bullet("Letting a non-technical teammate try a model.")
    d.bullet("Quickly iterating — dropping ten images in a row is faster than ten CLI invocations.")

    d.h3("pixo ui — full local dashboard")
    d.bold_line("Why it exists")
    d.p("Different purpose from pixo serve. Where pixo serve is a single-model "
        "interface, pixo ui is a control panel for everything — browse the "
        "model catalog, see past jobs, check hardware, set up cloud backends, "
        "manage environments.")
    d.bold_line("What it does")
    d.p("Launches a FastAPI + React dashboard on http://localhost:8420 with "
        "multiple pages: model catalog with privacy badges, job history, "
        "hardware profile, cloud backend setup.")
    d.bold_line("How to use it")
    d.code_block("pip install pixo[web]\npixo ui")

    d.h3("Which one should you use?")
    d.bullet("\"I want to run one model on some images\" → pixo serve <model>")
    d.bullet("\"I want to browse what pixo can do and check on past jobs\" → pixo ui")
    d.hr()

    # ---------- 9 ----------
    d.h2("9. Resource safety — not freezing your laptop")
    d.bold_line("Why it exists")
    d.p("Running a 4 GB SAM2 model on a laptop with 6 GB RAM will lock up the "
        "machine. Most tools don't warn you — they just start swapping, and "
        "your system becomes unresponsive. pixo's Resource Guardian prevents this.")

    d.h3("Pre-run safety check")
    d.p("Before any pixo run, pixo checks:")
    d.bullet("Free RAM vs model size (plus overhead estimate).")
    d.bullet("Free disk vs expected output size.")
    d.bullet("GPU VRAM if you're using --device cuda.")
    d.bullet("CPU temperature (warns if already hot).")
    d.p("If any of these look unsafe, you see a clear message and suggestions "
        "(smaller variant, --low-memory, cloud backend). Override with --force.")

    d.h3("Runtime limits")
    d.bullet("--max-ram <percent> — cap RAM usage (default: 70%).")
    d.bullet("--max-cpu <N> — limit to N CPU cores.")
    d.bullet("--background — runs at lowest OS priority. Your laptop stays fully usable during processing.")
    d.bullet("--low-memory — processes frame-by-frame with aggressive GC. Slower but works on 4 GB machines.")
    d.bold_line("Example")
    d.code_block(
        "pixo run yolov8 --input video.mp4                                  # standard\n"
        "pixo run yolov8 --input video.mp4 --low-memory --background --max-ram 40"
    )
    d.hr()

    # ---------- 10 ----------
    d.h2("10. Checkpointing — never lose progress")
    d.bold_line("Why it exists")
    d.p("Long video processing can take hours. If your laptop sleeps, battery "
        "dies, or you accidentally close the terminal — most tools make you "
        "start over. pixo saves progress automatically and picks up where it "
        "left off.")
    d.bold_line("How it works")
    d.bullet("Every N frames (default 100, configurable per model), pixo saves a checkpoint to ~/.pixo/jobs/<job_id>/.")
    d.bullet("Ctrl+C triggers a graceful pause instead of killing the process.")
    d.bullet("Ctrl+C a second time within 3 seconds actually force-quits.")
    d.bullet("Re-running the same command detects the checkpoint and asks if you want to resume.")
    d.bold_line("Commands")
    d.code_block(
        "pixo history            # see all past jobs\n"
        "pixo resume             # resume the most recent paused job\n"
        "pixo resume 3a37cea3    # resume a specific job\n"
        "pixo view 3a37cea3      # open output folder\n"
        "pixo jobs-clean         # delete completed jobs"
    )
    d.bold_line("Example flow")
    d.code_block(
        "pixo run yolov8 --input long_video.mp4\n"
        "# [Ctrl+C]\n"
        "# \"Paused at frame 1200/5000 (24%). Run pixo resume to continue.\"\n"
        "pixo resume\n"
        "# Picks up at frame 1201"
    )
    d.hr()

    # ---------- 11 ----------
    d.h2("11. Standard output structure")
    d.bold_line("Why it exists")
    d.p("Every model outputs different things (boxes, masks, depth maps, "
        "captions, OCR text). Without a standard structure, users are always "
        "hunting for \"where did the output go this time?\" pixo solves this "
        "by producing the same folder layout for every model.")
    d.bold_line("Layout")
    d.code_block(
        "pixo_output/\n"
        "  <jobid>_<model>_<inputname>/\n"
        "    results.json        # machine-readable metadata\n"
        "    summary.txt         # human-readable summary\n"
        "    visualizations/     # annotated images or video\n"
        "    raw/                # raw model output files\n"
        "    exports/            # COCO, CSV"
    )
    d.bold_line("Exports")
    d.bullet("COCO JSON — annotations in the COCO format that FiftyOne, Roboflow, and research frameworks understand.")
    d.bullet("CSV — flat tabular output for Excel or pandas.")
    d.p("Open the output folder with:")
    d.code_block("pixo view <job_id>")
    d.hr()

    # ---------- 12 ----------
    d.h2("12. Free cloud GPUs")
    d.bold_line("Why it exists")
    d.p("Some models (SAM2, SAMURAI, Grounding DINO, Florence-2 on video) are "
        "painfully slow on CPU. Instead of asking users to pay for AWS, pixo "
        "wires up two free options that anyone can use.")

    d.h3("Kaggle backend")
    d.p("Kaggle gives every account 30 hours per week of free T4 GPU.")
    d.bold_line("Setup (one time)")
    d.code_block(
        "pixo setup-cloud --kaggle\n"
        "# Enter Kaggle username and API key (from kaggle.com/settings)"
    )
    d.bold_line("Usage")
    d.code_block("pixo run sam2 --input photo.jpg --backend kaggle")
    d.p("pixo bundles your input, uploads it to Kaggle, runs the job on their "
        "GPU, and downloads the results. Typical SAM2 run: 1 minute on Kaggle "
        "vs 44 minutes on CPU.")

    d.h3("Colab backend")
    d.p("Google Colab gives free T4 GPU sessions.")
    d.code_block("pixo setup-cloud --colab\npixo run sam2 --input photo.jpg --backend colab")
    d.p("Colab integration is semi-automated — pixo generates a notebook and "
        "gives you step-by-step instructions. Kaggle is more seamless.")

    d.h3("Smart router")
    d.p("When you don't specify a backend, pixo estimates how long the run will "
        "take on each available backend and suggests the fastest one:")
    d.code_block(
        "Estimated times:\n"
        "  local:   7s      Available\n"
        "  kaggle:  1m 3s   Available\n"
        "  colab:   --      Not configured\n\n"
        "Recommended: local"
    )

    d.h3("Check status")
    d.code_block("pixo cloud-status")
    d.hr()

    # ---------- 13 ----------
    d.h2("13. ONNX optimization")
    d.bold_line("Why it exists")
    d.p("PyTorch models are convenient but not the fastest format for CPU "
        "inference. Converting to ONNX can make CPU inference 30–40% faster "
        "with no accuracy loss.")
    d.bold_line("What it does")
    d.p("pixo optimize <model> converts a downloaded PyTorch model to ONNX "
        "format. Next time you run the model, pixo detects the ONNX version "
        "and uses it automatically.")
    d.bold_line("How to use")
    d.code_block(
        "pixo pull yolov8\n"
        "pixo optimize yolov8\n"
        "pixo run yolov8 --input photo.jpg    # now ~40% faster on CPU"
    )
    d.p("The ONNX file sits next to the original weights in ~/.pixo/models/. "
        "Nothing else changes.")
    d.hr()

    # ---------- 14 ----------
    d.h2("14. Isolated environments")
    d.bold_line("Why it exists")
    d.p("Different CV models want different versions of torch, transformers, "
        "opencv. Installing them all together creates dependency conflicts. "
        "pixo solves this by letting each model live in its own virtual environment.")
    d.bold_line("How it works")
    d.code_block(
        "pixo pull sam2 --isolate        # creates ~/.pixo/envs/sam2/\n"
        "pixo run sam2 --input photo.jpg --isolate"
    )
    d.p("Each --isolate environment is a full Python venv with only that "
        "model's required packages. No conflicts with your base Python "
        "environment, no conflicts between models.")
    d.bold_line("Management")
    d.code_block(
        "pixo env-list              # show all isolated envs and sizes\n"
        "pixo env-clean sam2        # delete and rebuild sam2's env"
    )
    d.bold_line("When to use it")
    d.bullet("You're getting dependency errors running a model.")
    d.bullet("You want to keep your base Python environment clean.")
    d.bullet("You're running multiple different-generation models in the same session.")
    d.hr()

    # ---------- 15 ----------
    d.h2("15. Model piping")
    d.bold_line("Why it exists")
    d.p("Real CV workflows rarely use one model. You detect with YOLO, then "
        "segment the regions with SAM2. You find objects with Grounding DINO, "
        "then track them with SAMURAI. Doing this manually means running two "
        "commands, finding the output of the first, feeding it to the second "
        "— error-prone and tedious.")
    d.bold_line("What it does")
    d.p("Chains models together in one command. pixo handles converting outputs "
        "from one model into inputs for the next.")
    d.bold_line("How to use")
    d.code_block(
        "# Arrow syntax — explicit\n"
        "pixo pipe \"yolov8 -> depth_anything_v2\" --input photo.jpg\n"
        "pixo pipe \"grounding_dino -> sam2\" --input photo.jpg --prompt \"person\"\n\n"
        "# Pre-built templates\n"
        "pixo pipe detect_and_segment --input photo.jpg --prompt \"car\"\n"
        "pixo pipe segment_and_depth --input photo.jpg"
    )
    d.p("Each model in the pipeline produces its own output folder, and the "
        "final folder contains the combined result.")
    d.hr()

    # ---------- 16 ----------
    d.h2("16. Python SDK")
    d.bold_line("Why it exists")
    d.p("Not everyone wants to use the CLI. Data scientists, researchers, and "
        "anyone writing a Jupyter notebook wants a Python API.")
    d.bold_line("What's available")
    d.code_block(
        "import pixo\n\n"
        "# List all available models\n"
        "models = pixo.list_models()\n\n"
        "# Download a model\n"
        "pixo.pull(\"yolov8\")\n\n"
        "# Run inference\n"
        "result = pixo.run(\"yolov8\", input=\"photo.jpg\")\n"
        "print(result.objects)       # 12\n"
        "print(result.classes)       # ['person', 'car', 'chair']\n"
        "print(result.output_dir)    # './pixo_output/abc123_yolov8_photo'\n"
        "print(result.time_seconds)  # 7.2\n\n"
        "# Run with options\n"
        "result = pixo.run(\"grounding_dino\", input=\"photo.jpg\", prompt=\"person, car\")\n"
        "result = pixo.run(\"sam2\", input=\"photo.jpg\", backend=\"kaggle\")\n\n"
        "# Check hardware\n"
        "hw = pixo.doctor()\n"
        "print(hw[\"ram_total_gb\"], hw[\"has_gpu\"])\n\n"
        "# Pipeline\n"
        "result = pixo.pipe([\"yolov8\", \"depth_anything_v2\"], input=\"photo.jpg\")"
    )
    d.p("Everything the CLI does is available as a Python function, returning "
        "structured dataclass objects instead of terminal output.")
    d.hr()

    # ---------- 17 ----------
    d.h2("17. Hardware profiling")
    d.h3("pixo doctor")
    d.bold_line("Why it exists")
    d.p("Before running anything, users want to know: \"What can my machine "
        "actually handle?\"")
    d.bold_line("What it does")
    d.p("Shows a report with CPU name and cores, RAM, GPU/VRAM/CUDA, free "
        "disk space, CPU temperature, and a recommendation tailored to your "
        "hardware.")
    d.bold_line("Example output")
    d.code_block(
        "CPU:      Intel Core i7-1165G7 (4 cores)\n"
        "RAM:      16.0 GB total, 9.2 GB available\n"
        "GPU:      Not detected\n"
        "OS:       Windows 10.0\n"
        "Disk:     182.3 GB free\n"
        "CPU Temp: 52C (normal)\n\n"
        "Recommendation: Use ONNX + INT8 quantization for best CPU speed.\n"
        "                Run: pixo optimize <model>"
    )
    d.p("Run it once when you first install pixo to see what it suggests.")
    d.hr()

    # ---------- 18 ----------
    d.h2("18. Everyday housekeeping")
    d.h3("pixo rm <model>")
    d.p("Delete a downloaded model to free disk space.")
    d.code_block("pixo rm sam2\npixo rm sam2:tiny    # specific variant")

    d.h3("pixo upgrade")
    d.p("Update pixo itself to the latest version.")
    d.code_block("pixo upgrade")

    d.h3("pixo guide")
    d.p("In-terminal help text with the most common commands and tips.")
    d.code_block("pixo guide")

    d.h3("pixo jobs-clean")
    d.p("Remove checkpoints from completed jobs to free disk.")
    d.code_block("pixo jobs-clean")
    d.hr()

    # Locations
    d.h2("File locations reference")
    d.p("pixo stores data under ~/.pixo/:")
    d.table(
        ["Path", "What's in it"],
        [
            ["~/.pixo/models/", "Downloaded model weights"],
            ["~/.pixo/envs/", "Isolated per-model Python environments"],
            ["~/.pixo/jobs/", "Job checkpoints (for resume)"],
            ["~/.pixo/shares/", "Self-contained HTML share bundles"],
            ["~/.pixo/samples/", "Sample images used by pixo try"],
            ["~/.pixo/config.yaml", "Cloud backend credentials"],
            ["~/.pixo/logs/", "Detailed run logs"],
        ],
        col_widths=[2.2, 4.2],
    )
    d.p("Run outputs go to ./pixo_output/ by default (or wherever you set with --output).")
    d.hr()

    # Installs
    d.h2("Optional installs")
    d.p("pixo uses pip extras to keep the base install small:")
    d.table(
        ["Extra", "What it adds", "Install"],
        [
            ["yolo", "Ultralytics + OpenCV for YOLO-family models", "pip install pixo[yolo]"],
            ["onnx", "ONNX Runtime for optimized CPU inference", "pip install pixo[onnx]"],
            ["cloud", "Kaggle CLI for cloud backend", "pip install pixo[cloud]"],
            ["vision", "Transformers + PyTorch for Grounding DINO / Florence-2 / SAM2", "pip install pixo[vision]"],
            ["web", "FastAPI + uvicorn for pixo ui", "pip install pixo[web]"],
            ["demo", "Gradio for pixo serve", "pip install pixo[demo]"],
            ["all", "Everything", "pip install pixo[all]"],
        ],
        col_widths=[1.0, 3.4, 2.2],
    )
    d.p("Start with pip install pixo[yolo] if you just want detection. Add others as needed.")
    d.hr()

    # Cheat sheet
    d.h2("Quick start cheat sheet")
    d.code_block(
        "# Install\n"
        "pip install pixo[yolo]\n\n"
        "# 60-second demo\n"
        "pixo try\n\n"
        "# See what's available\n"
        "pixo list\n"
        "pixo doctor\n\n"
        "# Run a model\n"
        "pixo pull yolov8\n"
        "pixo run yolov8 --input photo.jpg\n\n"
        "# Share the result\n"
        "pixo share\n\n"
        "# Compare two models\n"
        "pixo compare yolov8 yolov11 --input photo.jpg\n\n"
        "# Run with privacy guarantee\n"
        "pixo run yolov8 --input medical.jpg --airgap\n\n"
        "# Browser UI\n"
        "pip install pixo[demo]\n"
        "pixo serve yolov8"
    )
    d.hr()

    d.p("That's all of pixo v0.3.0. For the one-page summary, see "
        "pixo_features_summary.docx. For the strategic roadmap, see "
        "pixo_differentiation_roadmap.md. For what's planned next, see "
        "pixo_v0.3_todo.md.")

    d.save("pixo_features_detailed.docx")


# ====================================================================

if __name__ == "__main__":
    build_summary()
    build_detailed()
    print("Done.")
